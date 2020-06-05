import codecs

from docx import Document


if __name__ == "__main__":
    docx_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万.docx'
    output_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万.txt'
    document = Document(docx_file_name)
    with codecs.open(output_file_name, mode='w', encoding='utf8') as fw:
        paragraph_len = len(document.paragraphs)
        print('paragraph_len:{}'.format(paragraph_len))
        # 按照段落读取
        for paragraph_id, paragraph in enumerate(document.paragraphs):
            if paragraph_id % 10 == 0:
                print('{}/{}'.format(paragraph_id, paragraph_len))
            line = paragraph.text
            line = line.strip()
            if len(line) > 0:
                fw.write('style：{}\n'.format(paragraph.style.name))  # 样式名称
                fw.write('{}\n'.format(line))
