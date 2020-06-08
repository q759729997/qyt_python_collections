import codecs
import json

from docx import Document


if __name__ == "__main__":
    # docx_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万.docx'
    # output_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万_tables.json'
    docx_file_name = './temp/1. 冀北电网调度控制管理规程（发文版）.docx'
    output_file_name = './temp/1. 冀北电网调度控制管理规程（发文版）_tables.json'
    document = Document(docx_file_name)
    with codecs.open(output_file_name, mode='w', encoding='utf8') as fw:
        table_len = len(document.tables)
        print('table_len:{}'.format(table_len))
        for table_id, table in enumerate(document.tables):
            if table_id % 5 == 0:
                print('{}/{}'.format(table_id, table_len))
            formated_table = list()
            for row in table.rows:
                formated_row = list()
                for column in row.cells:
                    text = column.text
                    formated_row.append(text)
                formated_table.append(formated_row)
            fw.write('{}\n'.format(json.dumps(formated_table, ensure_ascii=False)))
            # if paragraph_id > 100:
            #     break
