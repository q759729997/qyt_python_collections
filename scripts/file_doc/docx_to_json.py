import codecs
import json

from docx import Document


if __name__ == "__main__":
    # docx_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万.docx'
    # output_file_name = './temp/42--2020冀北电网调度细则--古杨树启动细则修改（0427）v04--张南85万.json'
    docx_file_name = './temp/1. 冀北电网调度控制管理规程（发文版）.docx'
    output_file_name = './temp/1. 冀北电网调度控制管理规程（发文版）.json'
    document = Document(docx_file_name)
    with codecs.open(output_file_name, mode='w', encoding='utf8') as fw:
        paragraph_len = len(document.paragraphs)
        print('paragraph_len:{}'.format(paragraph_len))
        # 按照段落读取
        for paragraph_id, paragraph in enumerate(document.paragraphs):
            if paragraph_id % 10 == 0:
                print('{}/{}'.format(paragraph_id, paragraph_len))
            text = paragraph.text
            text = text.strip()
            if len(text) > 0:
                style_obj = paragraph.style
                font_obj = paragraph.style.font
                row_data = {
                    'id': paragraph_id,
                    'text': text,
                    'style': {'name': style_obj.name},
                    'font': {'name': font_obj.name, 'size': font_obj.size, 'bold': font_obj.bold}
                }
                fw.write('{}\n'.format(json.dumps(row_data, ensure_ascii=False)))
            # if paragraph_id > 100:
            #     break
